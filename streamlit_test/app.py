import streamlit as st
from docx import Document
import pandas as pd
import io
import zipfile
from datetime import timedelta
from copy import deepcopy
from docx.enum.section import WD_SECTION
from docx.shared import Pt
from docx.oxml.ns import qn
import os

st.title("Hi！这里可以生成催缴函/回执函")

# =============================
# 提供 Excel 模板下载
# =============================
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
with open(os.path.join(BASE_DIR, "催缴函-template.xlsx"), "rb") as f:
    st.download_button(
        "📥 下载 Excel 模板（催缴函-template.xlsx）",
        data=f,
        file_name="催缴函-template.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )

# =============================
# 上传 Excel
# =============================
excel_file = st.file_uploader("上传已填写的Excel模板", type="xlsx")

# 选择生成类型
doc_type = st.selectbox("请选择生成类型：", ["催缴函", "回执函"])

# 日期选择器
if doc_type == "催缴函":
    send_date = st.date_input("请选择发函日期")
    stop_date = st.date_input("请选择支付欠费截止日期")
    end_date = stop_date + timedelta(days=1)
else:
    receipt_date = st.date_input("请选择回执日期")

if excel_file:
    st.success("Excel 上传成功！")
    df = pd.read_excel(excel_file)

    # 选择生成模式
    mode = st.radio(
        "请选择生成方式：",
        ("每个集团单独生成一个 Word", "合并所有集团到一个 PDF（格式永远不串行）")
    )

    # ---------------------------
    # 替换占位符函数
    # ---------------------------
    def replace_placeholder(doc, placeholders: dict, font_name=None, font_size=None):
        for p in doc.paragraphs:
            for key, value in placeholders.items():
                if key in p.text:
                    for run in p.runs:
                        if key in run.text:
                            run.text = run.text.replace(key, str(value))
                            if font_name:
                                run.font.name = font_name
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                            if font_size:
                                run.font.size = Pt(font_size)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in placeholders.items():
                        if key in cell.text:
                            for p in cell.paragraphs:
                                for run in p.runs:
                                    if key in run.text:
                                        run.text = run.text.replace(key, str(value))
                                        if font_name:
                                            run.font.name = font_name
                                            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                                        if font_size:
                                            run.font.size = Pt(font_size)

    # ---------------------------
    # 点击生成按钮
    # ---------------------------
    if st.button("生成 Word / PDF"):
        TEMPLATE1_PATH = os.path.join(BASE_DIR, "template1.docx")
        TEMPLATE2_PATH = os.path.join(BASE_DIR, "template2.docx")
        TEMPLATE_PATH = TEMPLATE1_PATH if doc_type == "催缴函" else TEMPLATE2_PATH

        # =====================================================
        # 1️⃣ 每个集团单独生成 Word（ZIP 方式）
        # =====================================================
        if mode == "每个集团单独生成一个 Word":
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zipf:
                for _, row in df.iterrows():
                    doc = Document(TEMPLATE_PATH)

                    if doc_type == "催缴函":
                        placeholders = {
                            "{{集团名称}}": row["集团名称"],
                            "{{客户经理}}": row["客户经理"],
                            "{{客户经理手机号}}": row["客户经理手机号"],
                            "{{逾期欠费金额}}": row["逾期欠费金额"],
                            "{{违约金}}": row["违约金"],
                            "{{共计欠费}}": row["共计欠费"],
                            "{{发函日期}}": send_date.strftime("%Y年%m月%d日"),
                            "{{支付欠费截止日期}}": stop_date.strftime("%Y年%m月%d日"),
                            "{{终止业务日期}}": end_date.strftime("%Y年%m月%d日"),
                        }
                        replace_placeholder(doc, placeholders)
                    else:
                        placeholders = {
                            "{{集团名称}}": row["集团名称"],
                            "{{客户经理}}": row["客户经理"],
                            "{{客户经理手机号}}": row["客户经理手机号"],
                            "{{共计欠费}}": row["共计欠费"],
                            "{{回执日期}}": receipt_date.strftime("%Y年%m月%d日"),
                        }
                        replace_placeholder(doc, placeholders, font_name="宋体", font_size=13)

                    file_buffer = io.BytesIO()
                    doc.save(file_buffer)
                    file_buffer.seek(0)
                    filename = f"{doc_type}_{row['集团名称']}.docx"
                    zipf.writestr(filename, file_buffer.getvalue())

            zip_buffer.seek(0)
            st.success("生成成功！点击下载 ZIP 文件👇")
            st.download_button(
                f"下载全部 {doc_type} Word（ZIP）",
                data=zip_buffer,
                file_name=f"{doc_type}合集.zip",
                mime="application/zip",
            )

        # =====================================================
        # 2️⃣ PDF 合并模式（格式绝不会串）
        # =====================================================
        else:
            from tempfile import NamedTemporaryFile
            from docx2pdf import convert
            from PyPDF2 import PdfMerger

            pdf_files = []

            for _, row in df.iterrows():
                doc = Document(TEMPLATE_PATH)

                if doc_type == "催缴函":
                    placeholders = {
                        "{{集团名称}}": row["集团名称"],
                        "{{客户经理}}": row["客户经理"],
                        "{{客户经理手机号}}": row["客户经理手机号"],
                        "{{逾期欠费金额}}": row["逾期欠费金额"],
                        "{{违约金}}": row["违约金"],
                        "{{共计欠费}}": row["共计欠费"],
                        "{{发函日期}}": send_date.strftime("%Y年%m月%d日"),
                        "{{支付欠费截止日期}}": stop_date.strftime("%Y年%m月%d日"),
                        "{{终止业务日期}}": end_date.strftime("%Y年%m月%d日"),
                    }
                    replace_placeholder(doc, placeholders)
                else:
                    placeholders = {
                        "{{集团名称}}": row["集团名称"],
                        "{{客户经理}}": row["客户经理"],
                        "{{客户经理手机号}}": row["客户经理手机号"],
                        "{{共计欠费}}": row["共计欠费"],
                        "{{回执日期}}": receipt_date.strftime("%Y年%m月%d日"),
                    }
                    replace_placeholder(doc, placeholders, font_name="宋体", font_size=13)

                # 保存 word → 转 pdf
                with NamedTemporaryFile(delete=False, suffix=".docx") as tmp_word:
                    tmp_word_path = tmp_word.name
                    doc.save(tmp_word_path)

                pdf_path = tmp_word_path.replace(".docx", ".pdf")
                convert(tmp_word_path, pdf_path)
                pdf_files.append(pdf_path)

            # ⭐ 合并 PDF
            merger = PdfMerger()
            for pdf in pdf_files:
                merger.append(pdf)

            merged_pdf_path = os.path.join(BASE_DIR, f"合并{doc_type}.pdf")
            merger.write(merged_pdf_path)
            merger.close()

            with open(merged_pdf_path, "rb") as f:
                st.success(f"合并 {doc_type} PDF 生成成功！（格式不会串行）")
                st.download_button(
                    f"下载合并版 {doc_type}（PDF）",
                    data=f,
                    file_name=f"合并{doc_type}.pdf",
                    mime="application/pdf",
                )

# 下面是你的重命名功能（保持不动）
# ============================================================
# 批量重命名工具
# ============================================================

import streamlit as st
import pandas as pd
import io
import zipfile
import os  # 用于处理文件名和后缀

st.title("这里可以批量重命名")

# ==========================
# 1️⃣ 提供 Excel 模板下载
# ==========================
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
with open(os.path.join(BASE_DIR, "Rename_template.xlsx"), "rb") as f:
    st.download_button(
        "📥 下载Excel 模板（Rename_template.xlsx）",
        data=f,
        file_name="Rename_template.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )

st.markdown("Tips: 按新名顺序扫描，扫描设置使用自动命名为1、2、3……")

st.image(os.path.join(BASE_DIR, "example.png"))

# ==========================
# 2️⃣ 上传 Excel
# ==========================
excel_file = st.file_uploader("上传已填写的 Excel 模板", type="xlsx")

if excel_file:
    df = pd.read_excel(excel_file)
    st.success("Excel 上传成功！")
    
    if "文件原名" not in df.columns or "新名" not in df.columns:
        st.error("Excel 必须包含列：'文件原名' 和 '新名'")
    else:
        df["文件原名"] = df["文件原名"].astype(str).str.strip().str.lstrip("'")
        df["新名"] = df["新名"].astype(str).str.strip().str.lstrip("'")

        files_to_rename = st.file_uploader(
            "选择需要重命名的文件（可以多选）",
            accept_multiple_files=True
        )

        if files_to_rename:
            st.write("已选择文件：", [f.name for f in files_to_rename])

            if st.button("开始批量重命名"):
                zip_buffer = io.BytesIO()
                renamed_count = 0

                with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zipf:
                    for f in files_to_rename:
                        file_base, file_ext = os.path.splitext(f.name)
                        file_base = file_base.strip().lstrip("'")

                        match_row = df[df["文件原名"] == file_base]
                        if not match_row.empty:
                            new_base_name = str(match_row["新名"].values[0]).strip()
                            new_name = new_base_name + file_ext
                            zipf.writestr(new_name, f.getbuffer())
                            renamed_count += 1
                        else:
                            st.warning(f"文件 '{f.name}' 在 Excel 中没有找到对应新名")

                zip_buffer.seek(0)
                st.success(f"重命名完成，共 {renamed_count} 个文件被重命名")
                st.download_button(
                    "📥 下载重命名后的文件（ZIP）",
                    data=zip_buffer,
                    file_name="重命名后的文件.zip",
                    mime="application/zip"
                )
